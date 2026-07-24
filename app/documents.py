"""
매뉴얼/문서 파일에서 텍스트를 추출하는 모듈.
관리자 웹의 대량 업로드 기능에서 사용 (음성/영상 STT와는 별개 경로).
"""
import io
import json
from fastapi import HTTPException


def try_parse_json_entries(data: bytes) -> list[dict] | None:
    """업로드된 파일이 [{"title":..., "content":...}, ...] 형태의 JSON이면 그 목록을 반환.
    JSON이 아니거나 형태가 안 맞으면 None (일반 텍스트로 처리하도록)."""
    try:
        parsed = json.loads(data.decode("utf-8"))
    except Exception:
        return None
    if not isinstance(parsed, list) or not parsed:
        return None
    if not all(isinstance(item, dict) and "content" in item for item in parsed):
        return None
    return parsed


def extract_text(filename: str, data: bytes) -> str:
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()

    if ext == "pdf":
        return _extract_pdf(data)
    if ext == "docx":
        return _extract_docx(data)
    if ext in ("txt", "md", "json"):
        return data.decode("utf-8", errors="ignore")

    raise HTTPException(status_code=422, detail=f"지원하지 않는 파일 형식입니다: .{ext} (pdf, docx, txt, md만 가능)")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(pages).strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail="PDF에서 텍스트를 추출하지 못했습니다. 스캔본(이미지) PDF는 OCR이 별도로 필요합니다.",
        )
    return text


def _extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 표 안의 텍스트도 함께 추출
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    text = "\n".join(parts).strip()
    if not text:
        raise HTTPException(status_code=422, detail="문서에서 텍스트를 추출하지 못했습니다.")
    return text
